from flask import Flask, request, render_template, jsonify, session, send_from_directory
import spacy
import openai, io
import os
import uuid
import json
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from transformers import BertTokenizer, BertModel
import torch
from docx import Document
import PyPDF2
from textblob import TextBlob
from nltk.tokenize import word_tokenize
from collections import Counter
from nltk.corpus import stopwords
from dotenv import load_dotenv
import secrets

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)

# Set the secret key from environment variables
app.secret_key = os.getenv('SECRET_KEY') or secrets.token_hex(16)

# Set your OpenAI API key from environment variables
openai.api_key = os.getenv('OPENAI_API_KEY', '')

# Load the spaCy model for entity extraction
nlp = spacy.load("en_core_web_md")

# Load BERT for semantic analysis
tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
bert_model = BertModel.from_pretrained('bert-base-uncased')

# Ensure NLTK stopwords are downloaded
import nltk
nltk.download('stopwords')

def extract_text_from_docx(file_stream):
    doc = Document(file_stream)
    return '\n'.join([p.text for p in doc.paragraphs])

def extract_text_from_pdf(file_stream):
    reader = PyPDF2.PdfFileReader(file_stream)
    text = []
    for page in range(reader.numPages):
        text.append(reader.getPage(page).extract_text())
    return '\n'.join(text)

def preprocess_text(text):
    return ' '.join(text.split())

def extract_skills(text):
    doc = nlp(text)
    skills = set()
    for ent in doc.ents:
        if ent.label_ in ['ORG', 'PRODUCT', 'SKILL']:
            skills.add(ent.text.strip())
    return skills

def get_bert_embeddings(text):
    inputs = tokenizer(text, return_tensors='pt', truncation=True, padding=True, max_length=512)
    with torch.no_grad():
        outputs = bert_model(**inputs)
    return outputs.last_hidden_state.mean(dim=1).numpy()

def calculate_matching_score(resume_text, job_description):
    vectorizer = TfidfVectorizer(stop_words='english')
    vectors = vectorizer.fit_transform([resume_text, job_description])
    tfidf_similarity = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
    
    resume_embeddings = get_bert_embeddings(resume_text)
    job_embeddings = get_bert_embeddings(job_description)
    semantic_similarity = cosine_similarity(resume_embeddings, job_embeddings)[0][0]
    
    final_score = (tfidf_similarity + semantic_similarity) / 2.0
    
    return final_score * 100

def filter_relevant_skills(matching_tags, non_matching_tags):
    return non_matching_tags

def get_keyword_density(text):
    stop_words = set(stopwords.words('english'))
    tokens = word_tokenize(text.lower())
    tokens = [token for token in tokens if token.isalnum() and token not in stop_words]
    frequency = Counter(tokens)
    return frequency.most_common(8)

def analyze_job_description(text):
    blob = TextBlob(text)
    language_tone = "Casual" if blob.sentiment.polarity < 0 else "Formal"
    
    if "team" in text.lower() or "collaborate" in text.lower():
        emphasis_teamwork = "Teamwork"
    else:
        emphasis_teamwork = "Individual Contribution"
    
    social_responsibility_terms = ["sustainability", "diversity", "ethics"]
    social_responsibility = any(term in text.lower() for term in social_responsibility_terms)
    
    keyword_density = get_keyword_density(text)
    
    return {
        "language_tone": language_tone,
        "emphasis_teamwork": emphasis_teamwork,
        "social_responsibility": social_responsibility,
        "keyword_density": keyword_density
    }

def count_tag_occurrences(text, tags):
    stop_words = set(stopwords.words('english'))
    tokens = word_tokenize(text.lower())
    tokens = [token for token in tokens if token.isalnum() and token not in stop_words]
    token_counts = Counter(tokens)
    
    tag_occurrences = {}
    for tag in tags:
        tag_occurrences[tag] = token_counts.get(tag.lower(), 0)
    
    return tag_occurrences

def update_resume_with_tags_using_llm(resume_text, selected_tags):
    prompt = (
        f"Here is a resume:\n\n{resume_text}\n\n"
        f"Here are some skills to add:\n{', '.join(selected_tags)}\n\n"
        f"Please integrate these skills into the resume in appropriate sections. "
        f"Ensure that the resume remains ATS optimized and the content is clear and professional."
    )

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500,
        temperature=0.7
    )

    updated_resume_text = response.choices[0].message['content'].strip()
    return updated_resume_text

def get_skill_use_cases_using_llm(tags):
    prompt = (
        f"Here are some skills:\n{', '.join(tags)}\n\n"
        f"For each skill, provide a real-world use case or example in 2-3 lines."
    )

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a knowledgeable assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500,
        temperature=0.7
    )

    return response.choices[0].message['content'].strip()

def save_text_to_docx(text, path):
    doc = Document()
    paragraphs = text.split('\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    doc.save(path)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/evaluate', methods=['POST'])
def evaluate_resume():
    resume_file = request.files.get('resume_file')
    job_description = request.form.get('job_description')
    
    if resume_file and resume_file.filename:
        file_ext = resume_file.filename.rsplit('.', 1)[1].lower()
        if file_ext == 'pdf':
            resume_text = extract_text_from_pdf(resume_file.stream)
        elif file_ext == 'docx':
            resume_text = extract_text_from_docx(resume_file.stream)
        else:
            return jsonify({'error': 'Unsupported file type. Please upload a PDF or DOCX file.'}), 400
    else:
        return jsonify({'error': 'No resume file provided.'}), 400

    resume_text = preprocess_text(resume_text)
    job_description = preprocess_text(job_description)
    
    match_score = calculate_matching_score(resume_text, job_description)

    resume_skills = extract_skills(resume_text)
    job_skills = extract_skills(job_description)
    
    missing_skills = job_skills - resume_skills
    matched_skills = job_skills & resume_skills

    missing_skills = filter_relevant_skills(matched_skills, list(missing_skills))

    job_analysis = analyze_job_description(job_description)
    
    # Include keyword density in the response
    # keyword_density = job_analysis.get('keyword_density', [])
    keyword_density = []
    # Get occurrences of matching and non-matching tags
    matching_tag_occurrences = count_tag_occurrences(job_description, list(matched_skills))
    non_matching_tag_occurrences = count_tag_occurrences(job_description, list(missing_skills))

    # Convert occurrences to the required format and append to keyword_density
    for keyword, count in matching_tag_occurrences.items():
        keyword_density.append({"keyword": keyword, "density": count})
    
    for keyword, count in non_matching_tag_occurrences.items():
        keyword_density.append({"keyword": keyword, "density": count})

    session['unique_id'] = str(uuid.uuid4())
    session['resume_text'] = resume_text
    session['missing_skills'] = list(missing_skills)

    return jsonify({
        'match_score': match_score,
        'relevant_tags': list(matched_skills),
        'non_matching_tags': list(missing_skills),
        'job_analysis': {
            'language_tone': job_analysis['language_tone'],
            'emphasis': job_analysis['emphasis_teamwork'],
            'social_responsibility': job_analysis['social_responsibility'],
            'keyword_density': keyword_density  # Include keyword density here
            
        }
    })




@app.route('/generate-updated-resume', methods=['POST'])
def generate_updated_resume():
    try:
        print("inside generate")
        selected_tags = request.json.get('tags', [])
        resume_text = session.get('resume_text', '')
        
        if not resume_text.strip():
            return jsonify({'error': 'No resume text found.'}), 400
        
        updated_resume_text = update_resume_with_tags_using_llm(resume_text, selected_tags)
        
        output_file_path = f"static/updated_resume_{session['unique_id']}.docx"
        save_text_to_docx(updated_resume_text, output_file_path)
        
        return jsonify({'file_path': output_file_path})
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'error': 'An error occurred while generating the updated resume.'}), 500

@app.route('/static/<path:filename>')
def download_file(filename):
    return send_from_directory('static', filename)

if __name__ == '__main__':
    app.run(debug=True)
