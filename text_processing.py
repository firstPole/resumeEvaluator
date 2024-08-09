import spacy
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
import PyPDF2
from nltk.tokenize import word_tokenize
from collections import Counter
from nltk.corpus import stopwords
import openai
import re
from textblob import TextBlob

# Load the spaCy model for entity extraction and embeddings
nlp = spacy.load("en_core_web_md")

# Ensure NLTK stopwords are downloaded
import nltk
nltk.download('stopwords')

def extract_text_from_docx(file_stream):
    doc = Document(file_stream)
    return '\n'.join([p.text for p in doc.paragraphs])

def extract_text_from_pdf(file_stream):
    try:
        reader = PyPDF2.PdfReader(file_stream)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        return None

def preprocess_text(text):
    return ' '.join(text.split())

def extract_skills(text):
    doc = nlp(text)
    skills = set()
    for ent in doc.ents:
        if ent.label_ in ['ORG', 'PRODUCT', 'WORK_OF_ART', 'EVENT']:  # Add more labels as needed
            skills.add(ent.text.lower())
    return skills

def get_spacy_embeddings(text):
    doc = nlp(text)
    return doc.vector

def calculate_jaccard_similarity(text1, text2):
    set1 = set(text1.split())
    set2 = set(text2.split())
    intersection = len(set1.intersection(set2))
    union = len(set1.union(set2))
    return intersection / union if union != 0 else 0

def calculate_matching_score(resume_text, job_description):
    # Extract skills from resume and job description
    resume_skills = extract_skills(resume_text)
    job_skills = extract_skills(job_description)

    # Calculate skills overlap
    skills_overlap = len(resume_skills & job_skills) / len(job_skills) if job_skills else 0

    # TF-IDF Vectorization
    vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1, 3))
    vectors = vectorizer.fit_transform([resume_text, job_description])
    tfidf_similarity = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
    
    # spaCy Embeddings
    resume_embedding = get_spacy_embeddings(resume_text)
    job_embedding = get_spacy_embeddings(job_description)
    semantic_similarity = cosine_similarity([resume_embedding], [job_embedding])[0][0]
    
    # Additional Metrics
    jaccard_similarity = calculate_jaccard_similarity(resume_text, job_description)
    
    # Enhanced Thresholds
    tfidf_threshold = 0.4
    semantic_threshold = 0.5
    jaccard_threshold = 0.3
    
    # Calculate Final Score with Minimum Score Handling
    if tfidf_similarity < tfidf_threshold and semantic_similarity < semantic_threshold and jaccard_similarity < jaccard_threshold:
        final_score = 0.1  # Set a very low score if no significant match is found
    else:
        # Weighted Average Calculation
        final_score = (0.4 * tfidf_similarity + 0.4 * semantic_similarity + 0.2 * jaccard_similarity)
    
    # Adjust final score based on skills overlap
    if skills_overlap == 0:
        final_score *= 0.1  # Decrease significantly if no skills overlap
    
    # Ensure final score is between 10% and 100%
    final_score = max(final_score, 0.1)
    
    return final_score * 100

def filter_relevant_skills(matching_tags, non_matching_tags):
    return non_matching_tags

def get_keyword_density(text):
    stop_words = set(stopwords.words('english'))
    tokens = word_tokenize(text.lower())
    tokens = [token for token in tokens if token.isalnum() and token not in stop_words]
    frequency = Counter(tokens)
    return frequency.most_common(8)

def count_tag_occurrences(text, tags):
    stop_words = set(stopwords.words('english'))
    tokens = word_tokenize(text.lower())
    tokens = [token for token in tokens if token.isalnum() and token not in stop_words]
    token_counts = Counter(tokens)
    
    tag_occurrences = {}
    for tag in tags:
        tag_occurrences[tag] = token_counts.get(tag.lower(), 0)
    
    return tag_occurrences

def update_resume_with_tags_using_llm(resume_text, selected_tags):
    prompt = (
        f"Here is a resume:\n\n{resume_text}\n\n"
        f"Here are some skills to add:\n{', '.join(selected_tags)}\n\n"
        f"Please integrate these skills into the resume in appropriate sections. "
        f"Ensure that the resume remains ATS optimized and the content is clear and professional."
    )

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500,
        temperature=0.7
    )

    updated_resume_text = response.choices[0].message['content'].strip()
    return updated_resume_text

def get_skill_use_cases_using_llm(tags):
    prompt = (
        f"Here are some skills:\n{', '.join(tags)}\n\n"
        f"For each skill, provide a real-world use case or example in 2-3 lines."
    )

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a knowledgeable assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500,
        temperature=0.7
    )

    return response.choices[0].message['content'].strip()

def save_text_to_docx(text, path):
    doc = Document()
    paragraphs = text.split('\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    doc.save(path)

def generate_personalized_recommendations(resume_text, job_description, matched_skills, missing_skills):
    # Use LLM to generate personalized recommendations
    prompt = (
        f"Here is a resume:\n\n{resume_text}\n\n"
        f"Here is a job description:\n\n{job_description}\n\n"
        f"Matched skills: {', '.join(matched_skills)}\n"
        f"Missing skills: {', '.join(missing_skills)}\n\n"
        f"Provide personalized recommendations for the candidate to improve their resume and better match the job description."
    )

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are an expert career advisor."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500,
        temperature=0.7
    )

    recommendations = response.choices[0].message['content'].strip()
    return recommendations

def extract_email(text):
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else None

def extract_phone_number(text):
    phone_pattern = r'\+?\d[\d -]{8,12}\d'
    phones = re.findall(phone_pattern, text)
    return phones[0] if phones else None

def extract_visa_info(text):
    # Define patterns to match visa-related sentences
    visa_pattern = r'\b(visa|assistance|sponsorship|support)\b.*?[.?!]'
    visa_matches = re.findall(visa_pattern, text, re.IGNORECASE)
    
    if not visa_matches:
        return 'No Visa Information Found'
    
    # Analyze sentiment of each extracted sentence
    for sentence in visa_matches:
        blob = TextBlob(sentence)
        sentiment = blob.sentiment.polarity
        
        # Sentiment polarity:
        # -1 to -0.1: Negative
        #  0 to 0.1: Neutral
        #  0.1 to 1: Positive
        if sentiment < 0:
            return 'Visa Sponsorship Not Provided'
        elif sentiment > 0:
            return 'Visa Sponsorship Provided'
    
    # If sentiment is neutral
    return 'Visa Information Unclear'

# # Example usage
# resume_text = "Extracted text from resume"
# job_description = "Extracted text from job description"
# matched_skills = ["skill1", "skill2"]
# missing_skills = ["skill3", "skill4"]

# matching_score = calculate_matching_score(resume_text, job_description)
# keyword_density = get_keyword_density(resume_text)
# relevant_skills = filter_relevant_skills(matched_skills, missing_skills)
# updated_resume_text = update_resume_with_tags_using_llm(resume_text, matched_skills)
# skill_use_cases = get_skill_use_cases_using_llm(matched_skills)
# recommendations = generate_personalized_recommendations(resume_text, job_description, matched_skills, missing_skills)

# print(f"Matching Score: {matching_score}")
# print(f"Keyword Density: {keyword_density}")
# print(f"Relevant Skills: {relevant_skills}")
# print(f"Updated Resume Text: {updated_resume_text}")
# print(f"Skill Use Cases: {skill_use_cases}")
# print(f"Personalized Recommendations: {recommendations}")
